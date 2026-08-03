"""HTML → Word 文档转换器。

用法:
    python export_docx.py input.html output.docx
    python export_docx.py input.html output.docx --title "报告标题"
    python export_docx.py - < output.docx  # 从 stdin 读取
    echo "<h1>标题</h1><table>..." | python export_docx.py - output.docx
"""

import sys
import os
import io
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: 需要 python-docx 库，请执行: pip install python-docx")
    sys.exit(1)


def _strip_tags(s):
    return re.sub(r'<[^>]+>', '', s).replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')


def html_to_docx(html_content, title="分析报告"):
    """将 HTML 内容转换为 Word 文档。

    Args:
        html_content: HTML 字符串
        title: 文档标题

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # 标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── HTML 表格 → Word 表格 ──
    table_re = re.compile(r'<table[^>]*>(.*?)</table>', re.DOTALL)
    tr_re = re.compile(r'<tr[^>]*>(.*?)</tr>', re.DOTALL)
    td_re = re.compile(r'<t[dh][^>]*>(.*?)</t[dh]>', re.DOTALL)

    # 分段处理：先分割 HTML 块（img 是自闭合标签，单独作为一块）
    blocks = re.split(r'(<(?:h[1-4]|table|p|hr|ul|ol|li|pre|blockquote)[^>]*>.*?</(?:h[1-4]|table|p|hr|ul|ol|li|pre|blockquote)>|<img[^>]*/?>)',
                      html_content, flags=re.DOTALL)

    for block in blocks:
        if not block or not block.strip():
            continue

        # 图片（base64 data URL → 嵌入 Word）
        img_match = re.match(r'<img[^>]*src="data:image/(?:png|jpeg|jpg);base64,([^"]+)"[^>]*>', block.strip())
        if img_match:
            try:
                import base64 as _b64
                img_data = _b64.b64decode(img_match.group(1))
                pic = doc.add_picture(io.BytesIO(img_data), width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            continue

        # 代码块（chartjson 等）：不导出为纯文本，避免 JSON 原文进文档
        if re.match(r'<pre[^>]*>', block.strip()):
            continue

        # 表格
        table_match = table_re.search(block)
        if table_match:
            rows_data = []
            for tr_match in tr_re.finditer(table_match.group(1)):
                cells = [_strip_tags(c) for c in td_re.findall(tr_match.group(1))]
                if cells:
                    rows_data.append(cells)

            if rows_data:
                cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=cols, style='Light Grid Accent 1')
                for ri, row in enumerate(rows_data):
                    for ci, cell_text in enumerate(row):
                        if ci < cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(9)
                doc.add_paragraph()
            continue

        # 标题
        m = re.match(r'<h(\d)[^>]*>(.*)</h\1>', block.strip(), re.DOTALL)
        if m:
            level = min(int(m.group(1)), 4)
            text = _strip_tags(m.group(2)).strip()
            if text:
                doc.add_heading(text, level=level)
            continue

        # 段落
        m = re.match(r'<(p|li|blockquote)[^>]*>(.*)</\1>', block.strip(), re.DOTALL)
        if m:
            text = _strip_tags(m.group(2)).strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(10)
            continue

        # 分隔线
        if re.match(r'<hr[^>]*>', block.strip()):
            doc.add_paragraph('─' * 40)
            continue

        # 纯文本（不在标签内的文本）
        text = _strip_tags(block).strip()
        if text and len(text) > 2:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    # 解析 --title 参数
    title = "QuantView 分析报告"
    for i, arg in enumerate(sys.argv):
        if arg == "--title" and i + 1 < len(sys.argv):
            title = sys.argv[i + 1]
            break

    # 读取输入
    if input_path == "-":
        html_content = sys.stdin.read()
    elif os.path.isfile(input_path):
        with open(input_path, "r", encoding="utf-8") as f:
            html_content = f.read()
    else:
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)

    if not html_content.strip():
        print("错误: 输入内容为空")
        sys.exit(1)

    # 转换
    buf = html_to_docx(html_content, title=title)
    with open(output_path, "wb") as f:
        f.write(buf.read())

    size_kb = os.path.getsize(output_path) / 1024
    print(f"已生成: {output_path} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
